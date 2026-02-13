from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_guide():
    doc = Document()

    # --- Header ---
    title = doc.add_heading('Navarasa Chitra: Interaction Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Introduction ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Navarasa Chitra is a collective emotional visualization project designed to transform individual human "
        "feelings into a growing, sacred geometric Mandala. Rooted in traditional Indian aesthetic "
        "theory, it maps the nine fundamental emotions (Navarasa) to ancient color systems (Sapta Varna) "
        "and the five primordial elements (Pancha Bhuta)."
    )

    # --- How to Use the Participant UI ---
    doc.add_heading('How to Interact (Participant UI)', level=1)
    doc.add_paragraph("Access the portal at: http://localhost:5001")

    # Step 1
    doc.add_heading('Step 1: Choose Your Rasa (The Lotus)', level=2)
    doc.add_paragraph(
        "Look at the central 9-petaled lotus. Each petal represents an emotion. Click a petal to begin designing that emotional state:"
    )
    rasas = [
        ("Shringara (Love)", "Vibrant and magnetic."),
        ("Hasya (Joy)", "Light and playful."),
        ("Karuna (Compassion)", "Tender and soft."),
        ("Raudra (Anger)", "Intense and piercing."),
        ("Veera (Heroism)", "Noble and powerful."),
        ("Bhayanaka (Fear)", "Ethereal and cautious."),
        ("Bibhatsa (Disgust)", "Deep and grounded."),
        ("Adbhuta (Wonder)", "Bright and expanding."),
        ("Shanta (Peace)", "Balanced and silent.")
    ]
    for r, desc in rasas:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{r}: ").bold = True
        p.add_run(desc)

    # Step 2
    doc.add_heading('Step 2: Bloom with Sapta Varna (Color & Transparency)', level=2)
    doc.add_paragraph(
        "Use the brightness/transparency sliders to imbue your chosen Rasa with color. In this project, 'Intensity' "
        "is interpreted as Transparency (Alpha). A 1% slider creates a ghostly, ethereal petal, while 100% creates a "
        "solid, bold ink stroke."
    )

    # Step 3
    doc.add_heading('Step 3: Imbue with Pancha Bhuta (Elemental Texture)', level=2)
    doc.add_paragraph(
        "Every emotion has a materiality. Select one of the five elements to apply a sacred texture to your petal:"
    )
    elements = [
        ("Earth (Prithvi)", "Solid, rectangular patterns indicating stability."),
        ("Water (Jala)", "Flowing wave patterns representing fluidity."),
        ("Fire (Agni)", "Sharp, triangular motifs for transformation."),
        ("Air (Vayu)", "Circular, swirling forms for movement."),
        ("Space (Akasha)", "Fine, stippled dots for expansion.")
    ]
    for e, desc in elements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{e}: ").bold = True
        p.add_run(desc)

    # Step 4
    doc.add_heading('Step 4: Contribute to the Whole', level=2)
    doc.add_paragraph(
        "Once your petal is ready, click 'Join the Collective'. You can configure multiple petals before submitting. "
        "Each submission adds a complete new 'Ring of Life' to the global mandala."
    )

    # --- The Global Mandala ---
    doc.add_heading('The Global Mandala Display', level=1)
    doc.add_paragraph("View the live artwork at: http://localhost:5001/mandala_display")
    doc.add_paragraph(
        "The global display uses a 'Semi-Made' geometric logic. As rings grow outward, the system maintains "
        "the constant size of the 'Onion Dome' petals (preserving their sacred shape) and multiplies them to fill the "
        "expanding circumference. This creates a dense, complex, and beautiful tapestry of many hearts beating as one."
    )

    # --- Setup Instructions ---
    doc.add_heading('Technical Setup', level=1)
    setup_p = doc.add_paragraph()
    setup_p.add_run("Directory: ").bold = True
    setup_p.add_run("/Users/shanmukhapilla/Downloads/VizChitra/Navarasa_Chitra\n")
    setup_p.add_run("Start Command: ").bold = True
    setup_p.add_run("source ../.venv/bin/activate && python app.py")

    # Save
    save_path = '/Users/shanmukhapilla/Downloads/VizChitra/Navarasa_Chitra/Navarasa_Chitra_User_Guide.docx'
    doc.save(save_path)
    print(f"Document saved to {save_path}")

if __name__ == "__main__":
    create_guide()
